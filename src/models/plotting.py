import plotly.express as px
import pandas as pd
from enum import Enum
from sklearn.metrics import accuracy_score, precision_score, recall_score, confusion_matrix, classification_report
from docx import Document
from pathlib import Path

class PlotMode(Enum):
    TEMPORAL = "temporal"
    SPATIAL = "spatial"
    BOTH = "both"

# Function to plot anomalies with Plotly
def plot_anomalies(
    parameter: list,
    df: pd.DataFrame,
    tn_sample_fraction: float = 0.1,
    mode: PlotMode = PlotMode.TEMPORAL,
    save_path: Path = None
):
    """
    Plots anomaly detection results for a given meteorological parameter using Plotly.

    Parameters:
    parameter: The name of the meteorological parameter to visualize (e.g., 'ETP', 'GLOT', 'RR', 'TN', 'TX').
    df: The dataset containing the original values, predicted anomalies, and real anomalies.
    mode: It allows to plot the temporal plots only, the spatial ones or both.
    The function creates a scatter plot with:
    - Blue points representing normal values.
    - Green points representing real anomalies (human corrections).
    - Red points representing predicted anomalies (detected by Isolation Forest).
    """
    test_data = df[df["is_test"] == 1]
    test_data["anomaly_type"] = test_data.apply(
        lambda row: "TP"
        if row[f"{parameter}_anomaly"] == 1 and row[f"{parameter}_anomaly_pred"] == 1
        else "FN"
        if row[f"{parameter}_anomaly"] == 1
        else "FP"
        if row[f"{parameter}_anomaly_pred"] == 1
        else "TN",
        axis=1,
    )
    # Downsample TN points
    tn_data = test_data[test_data["anomaly_type"] == "TN"].sample(
        frac=tn_sample_fraction, random_state=42
    )
    other_data = test_data[test_data["anomaly_type"] != "TN"]
    sampled_data = pd.concat([tn_data, other_data])
    if mode in [PlotMode.BOTH, PlotMode.TEMPORAL]:
        fig = px.scatter(
            sampled_data,
            x="datemesure",
            y=f"{parameter}_origine",
            color="anomaly_type",
            color_discrete_map={
                "TN": "blue",
                "FN": "green",
                "FP": "red",
                "TP": "purple",
            },
            title=f"Anomaly Detection for {parameter}",
            labels={"anomaly_type": "Anomaly Type"},
        )
        fig.update_layout(
            legend_title_text="Anomaly Type",
            legend=dict(
                title=None,
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1,
            ),
        )
        if save_path is not None:
            fig.write_image(save_path)
        else:
            fig.show()
    if mode in [PlotMode.BOTH, PlotMode.SPATIAL]:
        # fig = px.scatter_3d(sampled_data, x='Longitude', y='Latitude', z=parameter,
        #                     color='anomaly_type',
        #                     color_discrete_map={'TN': 'blue', 'FN': 'green', 'FP': 'red', 'TP': 'purple'},
        #                     title=f"3D Anomaly Detection for {parameter}",
        #                     labels={'anomaly_type': "Anomaly Type"},
        #                     hover_data=['Altitude'])
        # fig.update_traces(
        #     marker=dict(size=3, opacity=0.8)
        # )  # Adjust marker size & transparency
        fig = px.scatter(
            sampled_data,
            x="Longitude",
            y="Latitude",
            color="anomaly_type",
            color_discrete_map={
                "TN": "blue",
                "FN": "green",
                "FP": "red",
                "TP": "purple",
            },
            title=f"Anomaly Detection for {parameter}",
            labels={"anomaly_type": "Anomaly Type"},
            hover_data=["Altitude", parameter],
        )
        fig.update_yaxes(scaleanchor="x", scaleratio=1)
        if save_path is not None:
            fig.write_image(save_path)
        else:
            fig.show()


# Function to save results to a Word document
def save_results_to_word(df, parameters, modelname, filename="Anomaly_Report.doc", imagesdir=None):
    """
    Generates a Word document summarizing the anomaly detection results.

    The report includes:
    - A confusion matrix for each parameter, displaying actual vs. predicted anomalies.
    - Accuracy, precision, and recall metrics for each parameter.
    - Temporal and spatial anomaly plots for visualization.

    Parameters:
    ----------
    df : pd.DataFrame
        The dataframe containing anomaly detection results.
    parameters : list of str
        The list of meteorological parameters analyzed for anomalies.
    modelname : str
        Name of the model, used in figure filenames
    filename : str, optional
        The name of the output Word document (default is "Anomaly_Report.doc").
    
    Notes:
    ------
    - The function saves anomaly plots as temporary image files before inserting them into the document.
    - The temporary images are deleted after being added to the Word report.
    """
    doc = Document()
    doc.add_heading("Anomaly Detection Report", level=1)
    test_data = df[df["is_test"] == 1]
    
    for param in parameters:
        doc.add_heading(f"Results for {param}", level=2)
        cm = confusion_matrix(test_data[f"{param}_anomaly"], test_data[f"{param}_anomaly_pred"])
        acc = accuracy_score(test_data[f"{param}_anomaly"], test_data[f"{param}_anomaly_pred"])
        prec = precision_score(test_data[f"{param}_anomaly"], test_data[f"{param}_anomaly_pred"], zero_division=0)
        rec = recall_score(test_data[f"{param}_anomaly"], test_data[f"{param}_anomaly_pred"], zero_division=0)
        classif_report = classification_report(test_data[f"{param}_anomaly"], test_data[f"{param}_anomaly_pred"])
        
        doc.add_paragraph(f"Accuracy: {acc:.3f}, Precision: {prec:.3f}, Recall: {rec:.3f}")
        doc.add_paragraph(classif_report)
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        table.cell(0, 1).text = 'Predicted Normal'
        table.cell(0, 2).text = 'Predicted Anomaly'
        table.cell(1, 0).text = 'Actual Normal'
        table.cell(2, 0).text = 'Actual Anomaly'
        table.cell(1, 1).text = str(cm[0, 0])
        table.cell(1, 2).text = str(cm[0, 1])
        table.cell(2, 1).text = str(cm[1, 0])
        table.cell(2, 2).text = str(cm[1, 1])
        
        # Save plots and add them to the document
        plot_filename = f"{modelname}_{param}_anomaly_plot.png" if imagesdir is None else Path(imagesdir /  f"{modelname}_{param}_anomaly_plot.png")
        plotMode = PlotMode.BOTH if all(
            spatialcolumn in df.columns
            for spatialcolumn in ["Latitude", "Longitude", "Altitude", "cluster"]
        ) else PlotMode.TEMPORAL
        plot_anomalies(param, df, save_path=plot_filename, mode=plotMode)
        doc.add_picture(str(plot_filename))
    
    doc.save(filename)
